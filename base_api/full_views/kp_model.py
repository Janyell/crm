#!/usr/bin/python
# -*- coding: utf-8 -*-
from django.http import HttpResponseRedirect
from docx import Document
from docx.shared import Inches


def full_generate_kp(request):
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    document.add_picture('/Users/megge/Documents/Яндекс.Диск/gif/tumblr_nata787Rh01rjm4kxo1_500.jpg', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    row_cells = table.add_row().cells
    row_cells[0].text = str(3123)
    row_cells[1].text = str(3213)
    row_cells[2].text = str(2313213)

    document.add_page_break()

    document.save('demo.docx')

    # import os
    # from win32com import client
    #
    # wdFormatPDF = 17
    #
    # in_file = os.path.abspath('demo.docx')
    # out_file = os.path.abspath('demo.pdf')
    #
    # word = client.CreateObject('Word.Application')
    # doc = word.Documents.Open(in_file)
    # doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    # doc.Close()
    # word.Quit()
    return HttpResponseRedirect('/')